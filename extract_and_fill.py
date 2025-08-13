# extract_and_fill.py — fix24b (clean build)
import re
import unicodedata
from io import BytesIO
from typing import Dict, List, Tuple, Optional
from datetime import datetime
from zoneinfo import ZoneInfo

import pdfplumber
import pandas as pd
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_UNDERLINE
from docx.enum.table import WD_TABLE_ALIGNMENT

from docx.oxml import OxmlElement
from docx.oxml.ns import qn

DATE_RE = re.compile(r"\b([0-3]?\d)[./-]([01]?\d)[./-]([12]\d{3})\b")
COLUMNS_TARGET = ["Pos", "Référence", "Désignation", "Unité", "Qté", "Prix unit.", "Px u. Net", "Total CHF", "TVA"]

def today_ch() -> str:
    return datetime.now(ZoneInfo("Europe/Zurich")).strftime("%d.%m.%Y")

def _strip_accents(s: str) -> str:
    return "".join(c for c in unicodedata.normalize("NFKD", s) if not unicodedata.combining(c))

def _insert_missing_spaces(text: str) -> str:
    text = re.sub(r"(\d)(PC|PCE|KG|M|MM|CM|L)\b", r"\1 \2", text)
    text = re.sub(r"(\d)[A-Za-z]", lambda m: m.group(0)[0] + " " + m.group(0)[1:], text)
    return text

def extract_text_and_tables_from_pdf(file_like) -> Tuple[str, List[pd.DataFrame]]:
    texts = []
    tables: List[pd.DataFrame] = []
    with pdfplumber.open(file_like) as pdf:
        for page in pdf.pages:
            raw_text = page.extract_text() or ""
            raw_text = _insert_missing_spaces(raw_text)
            texts.append(raw_text)
            try:
                for raw in page.extract_tables() or []:
                    if not raw or len(raw) < 2:
                        continue
                    header = raw[0]; rows = raw[1:]
                    if not any(x for x in header): continue
                    df = pd.DataFrame(rows, columns=[(h or "").strip() for h in header])
                    if df.shape[1] >= 3 and df.shape[0] >= 1:
                        tables.append(_clean_df(df))
            except Exception:
                pass
            try:
                raw_single = page.extract_table()
                if raw_single and len(raw_single) > 1:
                    header = raw_single[0]; rows = raw_single[1:]
                    if any(x for x in header):
                        df = pd.DataFrame(rows, columns=[(h or "").strip() for h in header])
                        if df.shape[1] >= 3 and df.shape[0] >= 1:
                            tables.append(_clean_df(df))
            except Exception:
                pass
    unique, sigs = [], set()
    for df in tables:
        sig = (tuple(df.columns), df.shape)
        if sig not in sigs:
            sigs.add(sig); unique.append(df)
    return "\n".join(texts), unique

def _clean_df(df: pd.DataFrame) -> pd.DataFrame:
    df.columns = [str(c).strip() for c in df.columns]
    df = df.applymap(lambda x: (str(x).strip() if x is not None else ""))
    df = df.loc[~(df.apply(lambda r: all((not str(v).strip()) for v in r), axis=1))]
    return df.reset_index(drop=True)

def parse_fields_from_text(text: str) -> Dict[str, str]:
    """Use 'Total CHF' for the total displayed under the table; keep 'Montant Total TTC CHF' only for reference."""
    fields: Dict[str, str] = {}
    raw = text.replace("\xa0", " ")
    norm = _strip_accents(raw).lower()

    # Commande fournisseur (CF-..)
    m_norm = re.search(r"commande fournisseur n[°o]\s*([A-Za-z0-9\-_]+)", norm, flags=re.IGNORECASE)
    if m_norm:
        cf = m_norm.group(1).strip().upper()
        fields["N°commande fournisseur"] = cf
        fields["Commande fournisseur"] = cf

    # Notre référence (stop before any TVA token)
    m_line = re.search(r"(?i)(Notre\s+référence\s*:\s*)(.*)", raw)
    if m_line:
        after = m_line.group(2).strip()
        after_norm = _strip_accents(after).lower()
        cut_tokens = ["no tva", "n° tva", "n o tva", "tva", "no  tva"]
        cut_idx = None
        for tok in cut_tokens:
            idx = after_norm.find(tok)
            if idx != -1:
                cut_idx = idx; break
        value = after[:cut_idx].strip(" -–—\t·:;") if cut_idx is not None else after.strip(" -–—\t·:;")
        fields["Notre référence"] = value[:60]

    # Total CHF (used for display)
    total_chf = None
    m = re.search(r"(?i)Total\s+CHF\s*([0-9'’.,]+)", raw)
    if m:
        total_chf = m.group(1).strip()
    else:
        m = re.search(r"([0-9'’.,]+)\s*Total\s+CHF", raw, flags=re.IGNORECASE)
        if m:
            total_chf = m.group(1).strip()

    # Keep Montant Total TTC CHF for reference if present
    m_ttc = re.search(r"(?i)(Montant\s+Total\s+TTC\s+CHF|Total\s+TTC\s+CHF)\s*([0-9'’.,]+)", raw)
    if m_ttc:
        fields["Montant Total TTC CHF (PDF)"] = m_ttc.group(2).strip()
    else:
        m_ttc = re.search(r"([0-9'’.,]+)\s*(?i)(Montant\s+Total\s+TTC\s+CHF|Total\s+TTC\s+CHF)", raw)
        if m_ttc:
            fields["Montant Total TTC CHF (PDF)"] = m_ttc.group(1).strip()

    if total_chf:
        fields["Total TTC CHF"] = total_chf

    return fields

def _parse_date_str(s: str) -> Optional[datetime]:
    m = DATE_RE.search(s)
    if not m: return None
    d, mth, y = m.groups()
    try: return datetime(int(y), int(mth), int(d))
    except ValueError: return None

def latest_receipt_date_from_text(text: str) -> Optional[str]:
    dates = []
    norm = _strip_accents(text).lower().replace("\xa0", " ")
    lines = [l.strip() for l in norm.splitlines() if l.strip()]
    for ln in lines:
        if "delai de reception" in ln:
            dt = _parse_date_str(ln); 
            if dt: dates.append(dt)
    if not dates:
        for m in re.finditer(r"delai de reception.{0,30}?([0-3]?\d[./-][01]?\d[./-][12]\d{3})", norm):
            dt = _parse_date_str(m.group(0))
            if dt: dates.append(dt)
    if not dates: return None
    return max(dates).strftime("%d.%m.%Y")

def reconstruct_items_from_text(text: str) -> pd.DataFrame:
    """
    Start strictly at Pos multiples of 10; accumulate just enough until Total CHF is captured,
    then finalize the item immediately. Ignore interstitial meta lines.
    """
    unit_words = r"(PC|PCE|PCS|PIECE|PIECES|UN|UNITES?|KG|G|MG|L|ML|M|MM|CM)"
    money = r"[0-9'’.,]+"
    full_item_re = re.compile(
        rf"^\s*(?P<pos>\d{{1,4}})\s+"
        rf"(?P<ref>\d{{3,}})\s+"
        rf"(?P<designation>.+?)\s+"
        rf"(?P<unite>{unit_words})\s+"
        rf"(?P<qte>\d+)\s+"
        rf"(?P<pu>{money})\s+"
        rf"(?P<pxu>{money})\s+"
        rf"(?P<total>{money})"
        rf"(?:\s+(?P<tva>\d{{2,3}}))?\s*$",
        re.IGNORECASE
    )
    start_re = re.compile(r"^\s*(?P<pos>\d{1,4})\s+(?P<ref>\d{3,})\b")
    stop_cues = ("récapitulation", "recapitulation", "code tva", "montant total", "total ttc", "taux")
    junk_prefixes = ("tarif douanier", "pays d'origine", "indice :", "delai de reception :")

    rows: List[Dict[str, str]] = []
    buffer = ""
    buffering = False

    def try_flush(buf: str) -> bool:
        m = full_item_re.match(buf.strip())
        if not m: return False
        gd = m.groupdict()
        try:
            if int(gd["pos"]) % 10 != 0:
                return False
        except Exception:
            return False
        rows.append({
            "Pos": gd.get("pos",""),
            "Référence": gd.get("ref",""),
            "Désignation": (gd.get("designation","") or "").strip(),
            "Unité": gd.get("unite","").upper(),
            "Qté": gd.get("qte",""),
            "Prix unit.": gd.get("pu",""),
            "Px u. Net": gd.get("pxu",""),
            "Total CHF": gd.get("total",""),
            "TVA": gd.get("tva","") or "",
        })
        return True

    for raw_ln in text.splitlines():
        ln = raw_ln.strip()
        if not ln: continue
        low = _strip_accents(ln).lower()

        if any(low.startswith(p) for p in stop_cues):
            break
        if any(low.startswith(p) for p in junk_prefixes):
            continue

        if full_item_re.match(ln):
            try_flush(ln)
            buffering = False
            buffer = ""
            continue

        m_start = start_re.match(ln)
        if m_start:
            try:
                pos_val = int(m_start.group("pos"))
                if pos_val % 10 == 0:
                    buffering = True
                    buffer = ln
                    continue
                else:
                    continue
            except Exception:
                continue

        if buffering:
            if any(low.startswith(p) for p in junk_prefixes):
                continue
            buffer = (buffer + " " + ln).strip()
            if try_flush(buffer):
                buffering = False
                buffer = ""
                continue

    return pd.DataFrame(rows, columns=COLUMNS_TARGET)

def combine_detected_tables(tables: List[pd.DataFrame]) -> Optional[pd.DataFrame]:
    if not tables:
        return None
    candidate_rows = []
    for df in tables:
        dfc = _clean_df(df)
        cols_norm = [_strip_accents(str(c)).lower() for c in dfc.columns]
        if cols_norm and ("pos" not in cols_norm):
            first_col = dfc.columns[0]
            if dfc[first_col].astype(str).str.fullmatch(r"\d{1,4}").fillna(False).any():
                dfc = dfc.rename(columns={first_col: "Pos"})
        if "Pos" in dfc.columns:
            for _, r in dfc.iterrows():
                pos = str(r["Pos"]).strip()
                if pos.isdigit() and int(pos) % 10 == 0:
                    candidate_rows.append(r)
    if not candidate_rows:
        return None
    merged = pd.DataFrame(candidate_rows)
    for col in COLUMNS_TARGET:
        if col not in merged.columns:
            merged[col] = ""
    merged = merged[COLUMNS_TARGET]
    return merged.reset_index(drop=True)

def clean_items_df_keep_full(df: pd.DataFrame) -> pd.DataFrame:
    if df is None or df.empty: 
        return df
    filtered = []
    for _, r in df.iterrows():
        first_non_empty = ""
        for v in r.tolist():
            s = str(v).strip() if v is not None else ""
            if s: first_non_empty = _strip_accents(s.lower()); break
        if first_non_empty.startswith("indice :") or first_non_empty.startswith("delai de reception :"):
            continue
        filtered.append(r)
    new_df = pd.DataFrame(filtered, columns=df.columns) if filtered else df.iloc[0:0].copy()
    keep_cols = [c for c in new_df.columns if _strip_accents(str(c)).strip().lower() != "tva"]
    new_df = new_df[keep_cols] if keep_cols else new_df
    return new_df.reset_index(drop=True)

def replace_placeholders_everywhere(doc: Document, mapping: Dict[str, str]):
    def _replace_in_paragraph(paragraph, target: str, replacement: str):
        full_text = "".join(run.text for run in paragraph.runs)
        if target not in full_text: return
        new_text = full_text.replace(target, replacement)
        for idx in range(len(paragraph.runs)-1, -1, -1):
            r = paragraph.runs[idx]; r.clear(); r.text = ""
        if not paragraph.runs: paragraph.add_run(new_text)
        else: paragraph.runs[0].text = new_text
    for p in doc.paragraphs:
        for key, val in mapping.items():
            for variant in (f"« {key} »", f"«\xa0{key}\xa0»"):
                _replace_in_paragraph(p, variant, str(val))
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for key, val in mapping.items():
                        for variant in (f"« {key} »", f"«\xa0{key}\xa0»"):
                            _replace_in_paragraph(p, variant, str(val))
    for section in doc.sections:
        for hdr in (section.header, section.footer):
            for p in hdr.paragraphs:
                for key, val in mapping.items():
                    for variant in (f"« {key} »", f"«\xa0{key}\xa0»"):
                        _replace_in_paragraph(p, variant, str(val))

def compute_facture_suffix(fields: Dict[str, str]) -> Optional[str]:
    cmd = (fields.get("Commande fournisseur") or "").strip()
    m = re.search(r"CF-\d{2}-(\d+)", cmd, re.IGNORECASE)
    if m:
        return m.group(1)
    return None

def set_facture_title(doc: Document, suffix: Optional[str]):
    def set_para(p):
        txt = p.text.strip()
        if txt.startswith("Facture"):
            for r in p.runs[::-1]:
                r.clear()
            p.text = ""
            run = p.add_run(f"Facture {suffix}" if suffix else "Facture")
            run.bold = True
            run.font.size = Pt(12)
    for p in doc.paragraphs:
        set_para(p)
    for section in doc.sections:
        for hdr in (section.header, section.footer):
            for p in hdr.paragraphs:
                set_para(p)

def find_paragraph_anchor(doc: Document) -> Optional[object]:
    target_re = re.compile(r"cond\.\s*de\s*paiement[s]?", re.IGNORECASE)
    for p in doc.paragraphs:
        txt = _strip_accents(p.text).lower()
        if target_re.search(txt): return p
    return None

def _set_border(el, side, val='single', sz='8', space='0', color='auto'):
    border = el.find(qn(f'w:{side}'))
    if border is None:
        border = OxmlElement(f'w:{side}')
        el.append(border)
    border.set(qn('w:val'), val)
    border.set(qn('w:sz'), sz)
    border.set(qn('w:space'), space)
    border.set(qn('w:color'), color)

def set_table_borders_horizontal_only(table):
    """Outer frame + inside horizontals only; remove inside verticals."""
    tbl = table._element
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    tblBorders = tblPr.find(qn('w:tblBorders'))
    if tblBorders is None:
        tblBorders = OxmlElement('w:tblBorders'); tblPr.append(tblBorders)
    _set_border(tblBorders, 'top')
    _set_border(tblBorders, 'left')
    _set_border(tblBorders, 'bottom')
    _set_border(tblBorders, 'right')
    _set_border(tblBorders, 'insideH', val='single')
    _set_border(tblBorders, 'insideV', val='nil')
    tbl.append(tblPr)

    n_cols = len(table.rows[0].cells) if table.rows else 0
    for r_idx, row in enumerate(table.rows):
        for c_idx, cell in enumerate(row.cells):
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = tcPr.find(qn('w:tcBorders'))
            if tcBorders is None:
                tcBorders = OxmlElement('w:tcBorders'); tcPr.append(tcBorders)
            _set_border(tcBorders, 'top', val='single')
            _set_border(tcBorders, 'bottom', val='single')
            _set_border(tcBorders, 'left', val=('single' if c_idx == 0 else 'nil'))
            _set_border(tcBorders, 'right', val=('single' if c_idx == n_cols-1 else 'nil'))

def shade_header_row(table, fill_hex="EEF3FF"):
    if not table.rows: return
    for cell in table.rows[0].cells:
        tcPr = cell._tc.get_or_add_tcPr()
        shd = tcPr.find(qn('w:shd'))
        if shd is None:
            shd = OxmlElement('w:shd'); tcPr.append(shd)
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), fill_hex)

def apply_column_widths_and_alignments(table):
    try:
        table.autofit = False
    except Exception:
        pass
    header_cells = table.rows[0].cells
    headers = [c.text.strip() for c in header_cells]
    idx = {name: i for i, name in enumerate(headers)}
    widths_in = {}
    if "Pos" in idx: widths_in[idx["Pos"]] = Inches(0.5)
    if "Désignation" in idx: widths_in[idx["Désignation"]] = Inches(4.7)
    if "Qté" in idx: widths_in[idx["Qté"]] = Inches(0.8)
    for r in table.rows:
        for j, cell in enumerate(r.cells):
            if j in widths_in:
                for tcPr in cell._tc.iter(qn('w:tcPr')):
                    tcW = tcPr.find(qn('w:tcW'))
                    if tcW is None:
                        tcW = OxmlElement('w:tcW'); tcPr.append(tcW)
                    tcW.set(qn('w:type'), 'dxa')
                    dxa = int(widths_in[j].inches * 1440)
                    tcW.set(qn('w:w'), str(dxa))
            for p in cell.paragraphs:
                is_header = (r is table.rows[0])
                if is_header:
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                    if p.runs: p.runs[0].bold = True
                else:
                    if j in (idx.get("Pos", -1), idx.get("Référence", -1), idx.get("Désignation", -1)):
                        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                    elif j == idx.get("Qté", -1):
                        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    else:
                        if re.match(r"^\s*[0-9'’.,]+\s*$", p.text):
                            p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                        else:
                            p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

def insert_paragraph_after(paragraph, text=""):
    new_p = OxmlElement('w:p')
    paragraph._p.addnext(new_p)
    from docx.text.paragraph import Paragraph
    para = Paragraph(new_p, paragraph._parent)
    if text: para.add_run(text)
    return para

def insert_paragraph_after_element(elm, text="", align=None, bold=False, font_size_pt=None):
    new_p = OxmlElement('w:p')
    elm.addnext(new_p)
    from docx.text.paragraph import Paragraph
    para = Paragraph(new_p, elm.getparent())
    run = para.add_run(text)
    if bold: run.bold = True
    if font_size_pt: run.font.size = Pt(font_size_pt)
    if align is not None: para.alignment = align
    return para

def cleanup_extra_blank_paras(start_para, max_blank=2):
    blanks_kept = 0
    nxt = start_para._p.getnext()
    while nxt is not None and nxt.tag == qn('w:p'):
        texts = "".join(t.text or "" for t in nxt.iter(qn('w:t'))).strip()
        if texts == "":
            blanks_kept += 1
            if blanks_kept > max_blank:
                parent = nxt.getparent()
                to_remove = nxt
                nxt = nxt.getnext()
                parent.remove(to_remove)
                continue
        else:
            break
        nxt = nxt.getnext()

def add_total_row_to_table(table, label: str, amount: str):
    """Append a total row inside the table with merged label cells and double underline."""
    n_cols = len(table.rows[0].cells) if table.rows else 0
    if n_cols == 0:
        return
    row = table.add_row()
    cells = row.cells
    # Merge label cells (0..n-2) into cells[0]
    if n_cols >= 2:
        merged = cells[0]
        for j in range(1, n_cols-1):
            merged = merged.merge(cells[j])
        label_cell = merged
        amount_cell = cells[-1]
    else:
        label_cell = cells[0]
        amount_cell = cells[0]

    label_cell.text = label.strip()
    amount_text = (amount or "").strip()
    if amount_text:
        amount_cell.text = amount_text

    for c in (label_cell, amount_cell):
        for p in c.paragraphs:
            p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            if p.runs:
                p.runs[0].bold = True
                p.runs[0].font.underline = WD_UNDERLINE.DOUBLE

    # Add a strong double top border on the total row
    def _set_border(tcBorders, side, val='single', sz='18', color='auto'):
        el = tcBorders.find(qn(f'w:{side}'))
        if el is None:
            el = OxmlElement(f'w:{side}')
            tcBorders.append(el)
        el.set(qn('w:val'), val)
        el.set(qn('w:sz'), sz)
        el.set(qn('w:color'), color)
        el.set(qn('w:space'), '0')

    for c in row.cells:
        tc = c._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = tcPr.find(qn('w:tcBorders'))
        if tcBorders is None:
            tcBorders = OxmlElement('w:tcBorders'); tcPr.append(tcBorders)
        _set_border(tcBorders, 'top', val='double', sz='18')

def insert_df_two_lines_below_anchor(doc: Document, df: pd.DataFrame, total_ttc: Optional[str] = ""):
    if df is None or df.empty: return
    df = df.copy(); df.columns = [str(c) for c in df.columns]

    # Build table at end then move it near anchor
    tbl = doc.add_table(rows=1, cols=len(df.columns))
    hdr_cells = tbl.rows[0].cells
    for i, c in enumerate(df.columns):
        hdr_cells[i].text = str(c)
    for _, row in df.iterrows():
        cells = tbl.add_row().cells
        for i, col in enumerate(df.columns):
            val = "" if pd.isna(row[col]) else str(row[col])
            cells[i].text = val
            for para in cells[i].paragraphs:
                if para.runs: para.runs[0].font.size = Pt(10)

    # Styling
    shade_header_row(tbl, fill_hex="EEF3FF")
    set_table_borders_horizontal_only(tbl)
    apply_column_widths_and_alignments(tbl)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Move table 2 lines below anchor
    anchor = find_paragraph_anchor(doc)
    if anchor is not None:
        p1 = insert_paragraph_after(anchor, "")
        p2 = insert_paragraph_after(p1, "")
        p2._p.addnext(tbl._element)

    # Add total row inside the table
    add_total_row_to_table(tbl, label="Total TTC CHF", amount=total_ttc or "")

    # Exactly two blank lines after the table
    b1 = insert_paragraph_after_element(tbl._element, text="")
    b2 = insert_paragraph_after_element(b1._p, text="")

def process_pdf_to_docx(pdf_bytes: bytes, template_docx_bytes: bytes):
    text, tables = extract_text_and_tables_from_pdf(BytesIO(pdf_bytes))
    fields = parse_fields_from_text(text)
    fields["date du jour"] = today_ch()

    # Délai de réception max (alias Livré le)
    from_text = []
    norm = _strip_accents(text).lower().replace("\xa0"," ")
    for ln in [l.strip() for l in norm.splitlines() if l.strip()]:
        if "delai de reception" in ln:
            m = DATE_RE.search(ln)
            if m:
                d, mth, y = m.groups()
                try:
                    from_text.append(datetime(int(y), int(mth), int(d)))
                except Exception:
                    pass
    if from_text:
        max_dt = max(from_text).strftime("%d.%m.%Y")
        fields["Délai de réception"] = max_dt
        fields["Délai de livraison"] = max_dt

    # Items
    items_df = None
    if tables:
        items_df = combine_detected_tables(tables)
    if items_df is None:
        items_df = reconstruct_items_from_text(text)
    items_df = clean_items_df_keep_full(items_df)

    # Build doc with placeholders then title
    doc = Document(BytesIO(template_docx_bytes))
    replace_placeholders_everywhere(doc, fields)
    suffix = compute_facture_suffix(fields)
    set_facture_title(doc, suffix)

    out = BytesIO(); doc.save(out); out.seek(0)
    return out.getvalue(), fields, items_df

def build_final_doc(doc_bytes: bytes, items_df: pd.DataFrame, total_ttc: Optional[str]):
    doc = Document(BytesIO(doc_bytes))
    insert_df_two_lines_below_anchor(doc, items_df, total_ttc or "")
    out = BytesIO(); doc.save(out); out.seek(0)
    return out.getvalue()
