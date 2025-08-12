# streamlit_app.py — fix4: show "Délai de réception" and use it
import streamlit as st
import pandas as pd
from io import BytesIO
from docx import Document

from extract_and_fill import (
    process_pdf_to_docx,
    EXPECTED_COLUMNS,
    replace_placeholders_everywhere,
    find_or_create_items_table,
    insert_items_into_table,
    add_total_bottom_right,
)

st.set_page_config(page_title="PDF → DOCX (Commande fournisseur)", layout="wide")

st.title("PDF → DOCX : Remplissage automatique")
st.caption("Charge un PDF + un modèle DOCX. Corrige la détection si besoin, puis génère le document final.")

with st.sidebar:
    st.header("Étapes")
    st.markdown("1. Uploader le **PDF**")
    st.markdown("2. Uploader le **modèle .docx**")
    st.markdown("3. Vérifier/éditer le **tableau** et les **champs**")
    st.markdown("4. **Générer** et télécharger le `.docx`")

pdf_file = st.file_uploader("PDF de la commande", type=["pdf"])
docx_template = st.file_uploader("Modèle Word (.docx)", type=["docx"])

for k in ["detected_fields", "items_df", "tmpl_bytes", "pdf_bytes"]:
    if k not in st.session_state:
        st.session_state[k] = None

def _analyze():
    try:
        pdf_bytes = pdf_file.read()
        tmpl_bytes = docx_template.read()
        st.session_state["pdf_bytes"] = pdf_bytes
        st.session_state["tmpl_bytes"] = tmpl_bytes

        out_bytes, detected_fields, items_df = process_pdf_to_docx(
            pdf_bytes, tmpl_bytes, placeholder_overrides=None, custom_mapping=None
        )

        st.session_state["detected_fields"] = detected_fields or {}
        if items_df is None or items_df.empty:
            items_df = pd.DataFrame(columns=EXPECTED_COLUMNS)
        else:
            for col in EXPECTED_COLUMNS:
                if col not in items_df.columns:
                    items_df[col] = ""
            items_df = items_df[EXPECTED_COLUMNS]
        st.session_state["items_df"] = items_df

    except Exception as e:
        st.exception(e)

if pdf_file and docx_template and (st.session_state["detected_fields"] is None):
    with st.spinner("Analyse du PDF..."):
        _analyze()

if pdf_file and docx_template and st.session_state["detected_fields"] is not None:
    if st.button("🔁 Réanalyser à partir des fichiers envoyés"):
        with st.spinner("Ré-analyse du PDF..."):
            _analyze()

# Show placeholders list
if st.session_state["tmpl_bytes"]:
    with st.expander("Placeholders détectés dans le modèle"):
        doc = Document(BytesIO(st.session_state["tmpl_bytes"]))
        ph = []
        def collect(txt):
            if not txt:
                return
            parts = txt.split("«")
            for seg in parts[1:]:
                if "»" in seg:
                    ph.append(seg.split("»", 1)[0].strip(" «»\xa0"))
        for p in doc.paragraphs:
            collect(p.text)
        for t in doc.tables:
            for row in t.rows:
                for cell in row.cells:
                    collect(cell.text)
        st.write(sorted(set([p for p in ph if p])))

det = st.session_state["detected_fields"] or {}
st.subheader("Champs (auto-détectés, modifiables)")
col1, col2, col3 = st.columns(3)
with col1:
    f_cmd = st.text_input("N°commande fournisseur", value=det.get("N°commande fournisseur", ""))
    f_cmd2 = st.text_input("Commande fournisseur", value=det.get("Commande fournisseur", f_cmd))
    f_date = st.text_input("date du jour", value=det.get("date du jour", ""))
with col2:
    f_delai_recep = st.text_input("Délai de réception", value=det.get("Délai de réception", ""))
    f_cond = st.text_input("Cond. de paiement", value=det.get("Cond. de paiement", ""))
    f_total = st.text_input("Total TTC CHF", value=det.get("Total TTC CHF", ""))

# Map both keys so old templates still work
fields_over = {
    "N°commande fournisseur": f_cmd,
    "Commande fournisseur": f_cmd2,
    "date du jour": f_date,
    "Délai de réception": f_delai_recep,
    "date Délai de livraison": f_delai_recep,
    "Cond. de paiement": f_cond,
    "Total TTC CHF": f_total,
}

st.subheader("Tableau (éditable)")
items_df = st.session_state["items_df"]
if items_df is None:
    items_df = pd.DataFrame(columns=EXPECTED_COLUMNS)
edit_df = st.data_editor(items_df, num_rows="dynamic", use_container_width=True)

disabled = not (st.session_state["tmpl_bytes"] and st.session_state["pdf_bytes"])
if st.button("🧾 Générer le DOCX final", disabled=disabled):
    try:
        tmpl_bytes = st.session_state["tmpl_bytes"]
        if not tmpl_bytes:
            st.error("Template DOCX manquant.")
        else:
            doc = Document(BytesIO(tmpl_bytes))
            replace_placeholders_everywhere(doc, fields_over)
            table = find_or_create_items_table(doc, EXPECTED_COLUMNS)
            while len(table.rows) > 1:
                table._element.remove(table.rows[1]._element)
            insert_items_into_table(table, edit_df)
            add_total_bottom_right(doc, fields_over.get("Total TTC CHF", ""))

            out = BytesIO()
            doc.save(out); out.seek(0)
            st.success("DOCX généré !")
            st.download_button(
                "Télécharger le DOCX généré",
                data=out.getvalue(),
                file_name="commande_remplie.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
    except Exception as e:
        st.exception(e)
